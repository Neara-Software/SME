"""Convert markdown files to Word documents using python-docx."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def parse_markdown(text: str) -> list[dict]:
    """Parse markdown into a list of blocks."""
    blocks = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            content = line.lstrip("#").strip()
            blocks.append({"type": "heading", "level": level, "text": content})
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            i += 1
            continue

        # Code block
        if line.strip().startswith("```"):
            lang = line.strip().lstrip("`").strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append({"type": "code", "text": "\n".join(code_lines)})
            i += 1
            continue

        # Table
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|[\s\-:|]+\|\s*$", lines[i + 1]):
            rows = []
            while i < len(lines) and "|" in lines[i]:
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                # Skip separator row
                if not re.match(r"^[\s\-:|]+$", lines[i].strip().strip("|")):
                    rows.append(cells)
                i += 1
            blocks.append({"type": "table", "rows": rows})
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Paragraph (collect consecutive non-empty lines)
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and not lines[i].strip().startswith("```") and not ("|" in lines[i] and i + 1 < len(lines) and "|" in lines[i + 1]):
            para_lines.append(lines[i])
            i += 1
        if para_lines:
            blocks.append({"type": "paragraph", "text": " ".join(para_lines)})

    return blocks


def add_formatted_text(paragraph, text: str):
    """Add text with inline formatting (bold, code, italic)."""
    # Split on bold, code, and italic markers
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def convert_md_to_docx(md_path: Path, docx_path: Path):
    """Convert a markdown file to a Word document."""
    text = md_path.read_text(encoding="utf-8")
    blocks = parse_markdown(text)

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for block in blocks:
        if block["type"] == "heading":
            level = min(block["level"], 4)
            heading = doc.add_heading(block["text"], level=level)
            if level == 1:
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif block["type"] == "paragraph":
            para = doc.add_paragraph()
            add_formatted_text(para, block["text"])

        elif block["type"] == "code":
            para = doc.add_paragraph()
            para.style = doc.styles["Normal"]
            run = para.add_run(block["text"])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            para.paragraph_format.left_indent = Inches(0.5)

        elif block["type"] == "table":
            rows = block["rows"]
            if not rows:
                continue
            n_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=n_cols)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    if ci < n_cols:
                        cell = table.cell(ri, ci)
                        cell.text = ""
                        para = cell.paragraphs[0]
                        add_formatted_text(para, cell_text)
                        para.paragraph_format.space_after = Pt(2)
                        para.paragraph_format.space_before = Pt(2)
                        # Bold header row
                        if ri == 0:
                            for run in para.runs:
                                run.bold = True

            doc.add_paragraph()  # spacing after table

    doc.save(str(docx_path))
    print(f"  Created: {docx_path.name}")


def main():
    docs_dir = Path(__file__).parent
    md_files = sorted(docs_dir.glob("*.md"))

    if not md_files:
        print("No markdown files found!")
        sys.exit(1)

    print(f"Converting {len(md_files)} markdown files to Word...")
    for md_path in md_files:
        docx_path = md_path.with_suffix(".docx")
        convert_md_to_docx(md_path, docx_path)

    print("Done.")


if __name__ == "__main__":
    main()
